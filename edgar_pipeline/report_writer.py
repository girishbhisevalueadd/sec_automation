"""Word + PDF research report generator."""

from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

import config

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Section text parsing
# ---------------------------------------------------------------------------
SECTION_HEADERS = [
    "EXECUTIVE SUMMARY",
    "REVENUE TREND ANALYSIS",
    "MARGIN ANALYSIS",
    "BALANCE SHEET STRENGTH",
    "CASH FLOW QUALITY",
    "KEY RISKS",
    "OUTLOOK",
]


def _split_sections(narrative_text: str | None) -> dict[str, str]:
    sections = {h: "" for h in SECTION_HEADERS}
    if not narrative_text:
        return sections
    current: str | None = None
    buffer: list[str] = []
    for raw_line in narrative_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip().upper()
        matched = next((h for h in SECTION_HEADERS if stripped == h), None)
        if matched:
            if current:
                sections[current] = "\n".join(buffer).strip()
            current = matched
            buffer = []
            continue
        if current:
            buffer.append(line)
    if current:
        sections[current] = "\n".join(buffer).strip()
    return sections


# ---------------------------------------------------------------------------
# docx styling helpers
# ---------------------------------------------------------------------------
def _set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{config.REPORT_FOOTER} | Page ")
    run.font.size = Pt(9)
    # Page number field
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run2 = p.add_run()
    run2.font.size = Pt(9)
    run2._r.append(fld_begin)
    run2._r.append(instr)
    run2._r.append(fld_sep)
    run2._r.append(fld_end)
    p.add_run(" of ").font.size = Pt(9)
    # NUMPAGES field
    fld_begin2 = OxmlElement("w:fldChar")
    fld_begin2.set(qn("w:fldCharType"), "begin")
    instr2 = OxmlElement("w:instrText")
    instr2.text = " NUMPAGES "
    fld_sep2 = OxmlElement("w:fldChar")
    fld_sep2.set(qn("w:fldCharType"), "separate")
    fld_end2 = OxmlElement("w:fldChar")
    fld_end2.set(qn("w:fldCharType"), "end")
    run3 = p.add_run()
    run3.font.size = Pt(9)
    run3._r.append(fld_begin2)
    run3._r.append(instr2)
    run3._r.append(fld_sep2)
    run3._r.append(fld_end2)


def _df_to_table(doc: Document, df: pd.DataFrame, max_rows: int = 30) -> None:
    if df is None or df.empty:
        doc.add_paragraph("(no data available)").italic = True
        return
    work = df.copy()
    if "label" in work.columns:
        cols = ["label"] + [c for c in work.columns if c != "label"]
        work = work[cols]
    work = work.head(max_rows)
    n_rows, n_cols = work.shape
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    for j, col in enumerate(work.columns):
        cell = table.cell(0, j)
        cell.text = str(col)
        _set_cell_shading(cell, config.COLOR_HEADER_FILL)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
    # Data rows
    for i, (_, row) in enumerate(work.iterrows(), start=1):
        for j, col in enumerate(work.columns):
            v = row[col]
            if isinstance(v, float) and pd.notna(v):
                text = f"{v:,.0f}"
            elif pd.isna(v):
                text = ""
            else:
                text = str(v)
            cell = table.cell(i, j)
            cell.text = text
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
            if i % 2 == 0:
                _set_cell_shading(cell, config.COLOR_ALT_ROW_FILL)


def _section_heading(doc: Document, title: str, level: int = 1) -> None:
    h = doc.add_heading(title, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1D, 0x4E, 0x89)


def _load_filing_text(ticker: str, section: str) -> str:
    """Pull stored filing-text section (risk_factors / mda) for this ticker."""
    try:
        from storage import load_filing_text
        return load_filing_text(ticker, section, latest_only=True) or ""
    except Exception as e:  # noqa: BLE001
        logger.warning("filing text load failed (%s/%s): %s", ticker, section, e)
        return ""


def _split_text_for_doc(text: str, max_paragraphs: int = 60) -> list[str]:
    """Break a long text blob into reasonable paragraphs for python-docx.
    Caps at max_paragraphs to keep the Word file size manageable.
    """
    if not text:
        return []
    chunks = [p.strip() for p in text.split("\n\n") if p.strip()]
    if len(chunks) <= max_paragraphs:
        return chunks
    head = chunks[: max_paragraphs - 1]
    tail_note = f"[…truncated; {len(chunks) - len(head)} more paragraphs in source filing]"
    return head + [tail_note]


# ---------------------------------------------------------------------------
# Word builder
# ---------------------------------------------------------------------------
def build_word_report(
    ticker: str,
    summary_dict: dict[str, pd.DataFrame],
    ratios_df: pd.DataFrame,
    narrative_text: str | None = None,
    form_type: str = "10-K",
) -> Path:
    sections = _split_sections(narrative_text)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Cover page ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{ticker.upper()}")
    run.font.size = Pt(40)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1D, 0x4E, 0x89)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = subtitle.add_run("Financial Research Report")
    s_run.font.size = Pt(20)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(datetime.utcnow().strftime("%B %d, %Y")).font.size = Pt(12)

    prep_p = doc.add_paragraph()
    prep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prep_p.add_run(f"Prepared by {config.REPORT_PREPARED_BY}").italic = True

    doc.add_paragraph()
    logo_p = doc.add_paragraph("[Logo Placeholder]")
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ---- Table of Contents ----
    _section_heading(doc, "Table of Contents")
    toc_items = [
        "1. Company Overview",
        "2. Financial Summary",
        "3. Income Statement Analysis",
        "4. Balance Sheet Analysis",
        "5. Cash Flow Analysis",
        "6. Debt Profile (incl. Maturity Ladder)",
        "7. Segment & Geographic Detail",
        "8. Lease Commitments",
        "9. Stock-Based Compensation",
        "10. Income Tax Detail",
        "11. Risk Factors (from 10-K Item 1A)",
        "12. Management's Discussion & Analysis (from 10-K Item 7)",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    # ---- Section 1 ----
    _section_heading(doc, "1. Company Overview")
    overview = sections.get("EXECUTIVE SUMMARY") or (
        f"{ticker.upper()} is a US-listed public company. This report summarizes "
        "the most recent multi-period financials from its SEC filings."
    )
    doc.add_paragraph(overview)

    # ---- Section 2 - Financial Summary ----
    _section_heading(doc, "2. Financial Summary")
    doc.add_paragraph("Key ratios across recent fiscal periods:")
    if ratios_df is not None and not ratios_df.empty:
        _df_to_table(doc, ratios_df.reset_index())
    else:
        doc.add_paragraph("Ratio data unavailable.")

    # ---- Section 3 - Income Statement ----
    doc.add_page_break()
    _section_heading(doc, "3. Income Statement Analysis")
    _df_to_table(doc, summary_dict.get("income", pd.DataFrame()))
    doc.add_paragraph(sections.get("REVENUE TREND ANALYSIS") or
                       "[Revenue trend commentary placeholder.]")
    doc.add_paragraph(sections.get("MARGIN ANALYSIS") or
                       "[Margin commentary placeholder.]")

    # ---- Section 4 - Balance Sheet ----
    doc.add_page_break()
    _section_heading(doc, "4. Balance Sheet Analysis")
    _df_to_table(doc, summary_dict.get("balance", pd.DataFrame()))
    doc.add_paragraph(sections.get("BALANCE SHEET STRENGTH") or
                       "[Balance sheet commentary placeholder.]")

    # ---- Section 5 - Cash Flow ----
    doc.add_page_break()
    _section_heading(doc, "5. Cash Flow Analysis")
    _df_to_table(doc, summary_dict.get("cashflow", pd.DataFrame()))
    doc.add_paragraph(sections.get("CASH FLOW QUALITY") or
                       "[Cash flow commentary placeholder.]")

    # ---- Section 6 - Debt Profile ----
    doc.add_page_break()
    _section_heading(doc, "6. Debt Profile (incl. Maturity Ladder)")
    debt_df = summary_dict.get("debt", pd.DataFrame())
    maturity_df = summary_dict.get("debt_maturity", pd.DataFrame())
    if debt_df is not None and not debt_df.empty:
        _df_to_table(doc, debt_df)
    if maturity_df is not None and not maturity_df.empty:
        doc.add_paragraph().add_run("Maturity Ladder").bold = True
        _df_to_table(doc, maturity_df.reset_index() if "label" not in maturity_df.columns else maturity_df)
    if (debt_df is None or debt_df.empty) and (maturity_df is None or maturity_df.empty):
        doc.add_paragraph("No debt-specific XBRL facts were extracted from the latest filings.")

    # ---- Section 7 - Segment & Geography ----
    doc.add_page_break()
    _section_heading(doc, "7. Segment & Geographic Detail")
    seg_df = summary_dict.get("segment", pd.DataFrame())
    if seg_df is not None and not seg_df.empty:
        _df_to_table(doc, seg_df.reset_index() if "label" not in seg_df.columns else seg_df)
    else:
        doc.add_paragraph("Segment / geographic breakdown not separately reported in this filing.")

    # ---- Section 8 - Leases ----
    doc.add_page_break()
    _section_heading(doc, "8. Lease Commitments")
    lease_df = summary_dict.get("leases", pd.DataFrame())
    if lease_df is not None and not lease_df.empty:
        _df_to_table(doc, lease_df.reset_index() if "label" not in lease_df.columns else lease_df)
    else:
        doc.add_paragraph("Lease commitment detail not present in extracted facts.")

    # ---- Section 9 - SBC ----
    doc.add_page_break()
    _section_heading(doc, "9. Stock-Based Compensation")
    sbc_df = summary_dict.get("sbc", pd.DataFrame())
    if sbc_df is not None and not sbc_df.empty:
        _df_to_table(doc, sbc_df.reset_index() if "label" not in sbc_df.columns else sbc_df)
    else:
        doc.add_paragraph("Stock-based compensation detail not present in extracted facts.")

    # ---- Section 10 - Tax detail ----
    doc.add_page_break()
    _section_heading(doc, "10. Income Tax Detail")
    tax_df = summary_dict.get("tax_detail", pd.DataFrame())
    if tax_df is not None and not tax_df.empty:
        _df_to_table(doc, tax_df.reset_index() if "label" not in tax_df.columns else tax_df)
    else:
        doc.add_paragraph("Income tax detail not present in extracted facts.")

    # ---- Section 11 - Risk Factors (text from filing) ----
    doc.add_page_break()
    _section_heading(doc, "11. Risk Factors")
    risk_text = _load_filing_text(ticker, "risk_factors")
    if risk_text:
        for para in _split_text_for_doc(risk_text):
            doc.add_paragraph(para)
    else:
        doc.add_paragraph("Risk Factors text not available for this filing.")

    # ---- Section 12 - MD&A (text from filing) ----
    doc.add_page_break()
    _section_heading(doc, "12. Management's Discussion & Analysis")
    mda_text = _load_filing_text(ticker, "mda")
    if mda_text:
        for para in _split_text_for_doc(mda_text):
            doc.add_paragraph(para)
    else:
        doc.add_paragraph("MD&A text not available for this filing.")

    doc.add_page_break()
    doc.add_paragraph(sections.get("KEY RISKS") or "")
    doc.add_paragraph(sections.get("OUTLOOK") or "")

    _add_footer(doc)

    today = datetime.utcnow().strftime("%Y%m%d")
    form_safe = form_type.replace("-", "")
    out_path = config.REPORTS_DIR / f"{ticker.upper()}_{form_safe}_{today}.docx"
    doc.save(out_path)
    logger.info("Word report saved -> %s", out_path)
    return out_path


# ---------------------------------------------------------------------------
# PDF builder
# ---------------------------------------------------------------------------
def build_pdf_report(word_filepath: Path) -> Path:
    """Convert Word doc to PDF.

    Try docx2pdf first (uses MS Word on Windows / LibreOffice on Linux),
    fall back to a simple ReportLab-rendered PDF that re-reads the .docx.
    """
    word_filepath = Path(word_filepath)
    pdf_path = word_filepath.with_suffix(".pdf")

    try:
        from docx2pdf import convert  # type: ignore
        convert(str(word_filepath), str(pdf_path))
        if pdf_path.exists():
            logger.info("PDF saved via docx2pdf -> %s", pdf_path)
            return pdf_path
    except Exception as e:  # noqa: BLE001
        logger.warning("docx2pdf failed (%s) - trying ReportLab fallback.", e)

    # Fallback: re-read the docx text and lay it out with ReportLab.
    try:
        from docx import Document
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import (
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )
        from reportlab.lib import colors

        styles = getSampleStyleSheet()
        body_style = styles["BodyText"]
        h1 = styles["Heading1"]
        story = []

        src_doc = Document(word_filepath)
        for para in src_doc.paragraphs:
            text = (para.text or "").strip()
            if not text:
                story.append(Spacer(1, 6))
                continue
            style_name = (para.style.name or "").lower()
            if "heading" in style_name:
                story.append(Paragraph(text, h1))
            else:
                story.append(Paragraph(text, body_style))

        for tbl in src_doc.tables:
            data = []
            for row in tbl.rows:
                data.append([(cell.text or "").strip() for cell in row.cells])
            if not data:
                continue
            rl_table = Table(data, repeatRows=1)
            rl_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1D4E89")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
            ]))
            story.append(Spacer(1, 8))
            story.append(rl_table)
            story.append(Spacer(1, 8))

        rl_doc = SimpleDocTemplate(str(pdf_path), pagesize=LETTER)
        rl_doc.build(story)
        logger.info("PDF saved via ReportLab fallback -> %s", pdf_path)
        return pdf_path
    except Exception as e:  # noqa: BLE001
        logger.error("Both PDF backends failed: %s", e)
        raise
